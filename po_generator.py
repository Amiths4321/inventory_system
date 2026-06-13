# po_generator.py
import os
from pathlib  import Path
from datetime import datetime, timedelta

OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)


def generate_po_number() -> str:
    ts = datetime.now().strftime("%Y%m%d%H%M")
    return f"PO-{ts}"


def build_po_docx(
    recommendations: "pd.DataFrame",
    suppliers:       "pd.DataFrame" = None,
    company_name:    str = "TechCorp India Pvt Ltd"
) -> str:
    """Generate a formatted Purchase Order Word document."""
    from docx                   import Document
    from docx.shared            import Pt, RGBColor
    from docx.enum.text         import WD_ALIGN_PARAGRAPH

    doc      = Document()
    po_num   = generate_po_number()
    po_date  = datetime.now().strftime("%B %d, %Y")
    due_date = (datetime.now() + timedelta(days=14)).strftime("%B %d, %Y")

    # Header
    title = doc.add_heading("PURCHASE ORDER", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"PO Number: {po_num}")
    doc.add_paragraph(f"Date: {po_date}")
    doc.add_paragraph(f"Required By: {due_date}")
    doc.add_paragraph(f"Issued By: {company_name}")
    doc.add_paragraph()

    # Group by supplier
    if "Supplier" in recommendations.columns:
        grouped = recommendations.groupby("Supplier")
    else:
        grouped = [("General", recommendations)]

    total_po_value = 0

    for supplier_id, items in grouped:
        # Supplier info
        supplier_name = supplier_id
        if suppliers is not None and not suppliers.empty:
            sup_row = suppliers[suppliers["Supplier ID"] == supplier_id]
            if not sup_row.empty:
                supplier_name = sup_row.iloc[0].get("Supplier Name", supplier_id)

        doc.add_heading(f"Supplier: {supplier_name} ({supplier_id})", level=2)

        # Items table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"

        # Header row
        hdr    = table.rows[0].cells
        hdr[0].text = "SKU"
        hdr[1].text = "Product"
        hdr[2].text = "Qty"
        hdr[3].text = "Unit Price (Rs.)"
        hdr[4].text = "Total (Rs.)"

        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True

        # Data rows
        subtotal = 0
        for _, row in items.iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row.get("SKU", ""))
            cells[1].text = str(row.get("Product", ""))
            cells[2].text = str(row.get("Recommended Qty", 0))
            cells[3].text = f"{row.get('Unit Price', 0):,.2f}"
            total         = row.get("Total Value", 0)
            cells[4].text = f"{total:,.2f}"
            subtotal     += total

        total_po_value += subtotal
        doc.add_paragraph(f"Subtotal: Rs. {subtotal:,.2f}")
        doc.add_paragraph()

    # Total
    doc.add_heading(f"TOTAL PO VALUE: Rs. {total_po_value:,.2f}", level=2)

    doc.add_paragraph()
    doc.add_paragraph("Terms and Conditions:")
    doc.add_paragraph("1. Deliver to the address specified by the procurement team.")
    doc.add_paragraph("2. Invoice must reference this PO number.")
    doc.add_paragraph("3. Payment will be made within 30 days of delivery and invoice receipt.")
    doc.add_paragraph()
    doc.add_paragraph("Authorised By: _______________________")
    doc.add_paragraph("Signature: ___________________________")
    doc.add_paragraph(f"Date: {po_date}")

    filepath = OUTPUT_DIR / f"{po_num}.docx"
    doc.save(str(filepath))
    return str(filepath)